from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


BASE = Path(__file__).resolve().parent
ASSET = BASE / "report_assets"
OUT = BASE / "202426001054_吴凯珍_光学系统设计课程报告.docx"


def font_run(run, size=12, bold=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_cell_text(cell, text, bold=False, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1
    run = p.add_run(text)
    font_run(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.paragraph_format.line_spacing = 1
        style.paragraph_format.space_after = Pt(0)
    styles["Title"].font.size = Pt(18)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True

    def add_para(text="", align=None, bold=False, size=12, first_indent=True):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if first_indent and text:
            p.paragraph_format.first_line_indent = Pt(24)
        if align is not None:
            p.alignment = align
            p.paragraph_format.first_line_indent = None
        run = p.add_run(text)
        font_run(run, size=size, bold=bold)
        return p

    def add_heading(text, level=1):
        p = doc.add_paragraph(style=f"Heading {level}")
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_before = Pt(6 if level == 1 else 3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = None
        run = p.add_run(text)
        font_run(run, size=(14 if level == 1 else 13), bold=True)
        return p

    def add_caption(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        font_run(run, size=10.5)

    def add_picture(img_name, caption, width=5.7):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(ASSET / img_name), width=Inches(width))
        add_caption(caption)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("《光学系统设计》课程报告")
    font_run(r, size=20, bold=True)
    for _ in range(4):
        add_para("", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

    cover = doc.add_table(rows=4, cols=2)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(cover)
    cover_data = [
        ("题目", "基于库克三片式镜头的光学系统优化设计"),
        ("班级", "24级光电信息科学与工程"),
        ("姓名", "吴凯珍"),
        ("学号", "202426001054"),
    ]
    for row, (key, value) in zip(cover.rows, cover_data):
        row.height = Cm(1.1)
        set_cell_text(row.cells[0], key, bold=True, size=14)
        set_cell_text(row.cells[1], value, size=14)
    for _ in range(6):
        add_para("", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    add_para("完成日期：2026年6月28日", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

    doc.add_page_break()

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("基于库克三片式镜头的光学系统优化设计")
    font_run(r, size=18, bold=True)
    add_para(
        "班级：24级光电信息科学与工程    姓名：吴凯珍    学号：202426001054",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
    )

    add_heading("一、背景介绍与设计目标")
    add_para(
        "库克三片式物镜（Cooke Triplet）是一种经典的三片分离式成像镜头，通常由正透镜、负透镜和正透镜组成。"
        "该结构利用正负光焦度的合理分配，在较少镜片数量下同时校正球差、彗差、像散、场曲、畸变和轴向色差，"
        "因而常用于照相物镜、小型成像系统以及光学设计教学中的初始结构优化。"
    )
    add_para(
        "本次课程设计以 Zemax OpticStudio 18.9 中的库克三片式初始镜头为基础，按照题目给出的指标进行参数优化。"
        "设计目标为：有效焦距 EFFL = 10 mm，相对孔径为 1/2.8，半视场角为 15°（即 2ω = 30°），"
        "并要求在空间频率 100 lp/mm 处复色光衍射 MTF 大于 0.4。"
    )
    add_para(
        "该设计的核心任务不是重新建立复杂新结构，而是在已有三片式初始结构上选择合适变量、建立评价函数并进行优化，"
        "使系统在满足焦距、口径、视场和工艺边界的同时获得较好的成像质量。"
    )

    add_heading("二、Zemax 系统设置与初始结构")
    add_heading("1、系统孔径、视场与波长设置", 2)
    add_para(
        "系统孔径采用入瞳直径方式设置。根据 EFFL = 10 mm 和相对孔径 1/2.8，可得 F 数为 2.8，"
        "入瞳直径约为 10/2.8 = 3.571 mm，与截图中的系统孔径设置一致。视场设置为 0°、7.5°、15°三个视场，"
        "权重均为 1，用于兼顾轴上、半视场和边缘视场像质。波长采用可见光 F、d、C 三条谱线，分别约为 "
        "0.486 μm、0.588 μm、0.656 μm，权重均为 1，用于考察复色光成像性能。"
    )
    add_picture("step_02.png", "图1 系统孔径、视场与波长设置", width=3.4)

    add_heading("2、初始镜头结构", 2)
    add_para(
        "初始结构为 Zemax 示例库中的 A SIMPLE COOKE TRIPLET，采用三片分离透镜形式。由初始布局图可以看出，"
        "系统由前组正透镜、中间负透镜和后组正透镜组成，光阑位于中间附近，有利于平衡离轴像差。"
        "初始系统总轴长约为 42.85 mm，光线在像面附近已能成像，但点列图显示离轴视场弥散明显，"
        "仍需要通过曲率、厚度、空气间隔和像面位置等变量进一步优化。"
    )
    add_picture("step_04.png", "图2 初始三片式镜头二维布局图", width=5.2)
    add_picture("step_03.png", "图3 初始结构点列图", width=5.2)
    add_para(
        "初始点列图中，三个视场的 RMS 半径分别约为 6.273 μm、13.151 μm、17.040 μm，"
        "边缘视场弥散斑明显大于轴上视场，说明初始结构在较大视场下的彗差、像散和场曲尚未得到充分平衡。"
    )

    add_heading("三、优化设置与过程分析")
    add_heading("1、评价函数设置", 2)
    add_para(
        "优化评价函数采用以对比度为目标的默认评价函数，空间频率设置为 100 lp/mm，类型为 RMS，"
        "高斯求积采样为 3 环 6 臂。该设置直接对应题目中“100 lp/mm 处 MTF > 0.4”的要求，"
        "使优化过程围绕高空间频率成像对比度展开。"
    )
    add_para(
        "同时设置厚度边界条件：玻璃最小厚度 0.3 mm、最大厚度 15 mm，空气最小间隔 0.3 mm，边缘厚度 0.3 mm；"
        "并加入有效焦距 EFFL = 10 mm 的约束。这样可以避免优化只追求像质而产生过薄镜片、负空气间隔或焦距偏离目标等不可实现结果。"
    )
    add_picture("step_01.png", "图4 评价函数与边界约束设置", width=5.7)

    add_heading("2、变量选择与迭代优化", 2)
    add_para(
        "优化变量主要包括各透镜表面的曲率半径、镜片厚度、空气间隔以及像面位置等。第一轮优化从原始库克三片式结构出发，"
        "使用阻尼最小二乘法进行自动优化；随后根据布局和点列图结果继续调整变量范围，并在保持三片式基本结构不变的前提下改善边缘视场像质。"
    )
    add_para(
        "从截图可见，优化过程中评价函数由较大的初始值逐步降低。例如一次优化窗口显示初始评价函数约为 17.5033，"
        "优化后当前评价函数约为 0.0753，说明系统像差和约束误差得到显著改善。该过程体现了 Zemax 中"
        "“建立评价函数—设置变量—运行优化—检查图形结果—再优化”的典型设计流程。"
    )
    add_picture("step_12.png", "图5 阻尼最小二乘法优化窗口", width=5.7)

    add_heading("3、结构参数变化", 2)
    add_para(
        "优化后镜头仍保持正、负、正三片式形式，但各面的曲率和空气间隔发生了明显变化。初始结构的第一片材料为 SK16，"
        "中间片材料为 F2，第三片材料为 SK16；优化后仍沿用这一典型玻璃组合，用冠牌玻璃与火石玻璃配合来校正色差。"
    )
    add_para(
        "从最终镜头数据可以读出，优化后第 1 面曲率半径约为 22.014 mm，厚度约为 3.259 mm，材料为 SK16；"
        "第 2 面曲率半径约为 -435.760 mm，空气间隔约为 6.008 mm；第 3 面曲率半径约为 -22.213 mm，"
        "厚度约为 1.000 mm，材料为 F2；光阑位于第 4 面附近；第 5 面曲率半径约为 79.684 mm，厚度约为 2.952 mm，"
        "材料为 SK16；第 6 面曲率半径约为 -18.395 mm，至像面距离约为 42.208 mm。最终像面半径约为 13.357 mm。"
    )
    add_picture("step_13.png", "图6 最终优化后的镜头数据编辑器", width=5.7)
    add_picture("step_15.png", "图7 优化后系统布局图", width=5.7)

    add_heading("四、运行图形结果与像质分析")
    add_heading("1、点列图分析", 2)
    add_para(
        "点列图用于观察不同视场和不同波长光线在像面上的弥散情况。优化后系统在 0°、7.5°、15°三个视场下均能形成较集中的弥散斑，"
        "且三条波长的弥散分布没有出现严重分离，说明色差得到一定控制。"
    )
    add_para(
        "最终点列图显示，0°、7.5°、15°三个视场的 RMS 半径分别约为 7.927 μm、23.040 μm、12.881 μm。"
        "与初始结构相比，边缘视场的几何像差分布得到改善，尤其是离轴光线相对像面的集中程度提高。"
        "虽然中间视场仍存在一定弥散，但整体成像质量已能配合 MTF 指标满足课程设计要求。"
    )
    add_picture("step_10.png", "图8 优化过程中点列图结果", width=5.0)
    add_picture("step_14.png", "图9 优化后点列图结果", width=5.2)

    add_heading("2、MTF 分析", 2)
    add_para(
        "MTF 曲线是评价成像系统对不同空间频率细节传递能力的重要指标。本设计重点关注 100 lp/mm 处的复色光衍射 MTF。"
        "由截图可见，在 0°、7.5°、15°三个视场中，子午和弧矢方向的 MTF 曲线随空间频率升高逐渐下降，"
        "但在 100 lp/mm 处仍保持在约 0.55 以上，明显高于题目要求的 0.4。"
    )
    add_para(
        "这说明优化后的三片式镜头在目标空间频率处仍具有足够的对比度传递能力。轴上视场曲线最高，离轴视场略低，"
        "符合较大视场光学系统的一般规律；子午和弧矢曲线之间存在一定差异，反映了离轴像散和场曲仍未完全消除，"
        "但对课程目标而言已经达到合格像质。"
    )
    add_picture("step_05.png", "图10 优化后复色光衍射 MTF 曲线", width=5.7)

    add_heading("3、综合结果", 2)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ["项目", "设计要求", "优化结果"]):
        set_cell_text(cell, text, bold=True)
    rows = [
        ("有效焦距", "EFFL = 10 mm", "评价函数中加入 EFFL 约束，按 10 mm 目标优化"),
        ("相对孔径", "1/2.8", "入瞳直径约 3.571 mm，对应 F/# = 2.8"),
        ("视场", "2ω = 30°", "设置 0°、7.5°、15°三个视场"),
        ("波长", "可见光复色", "F、d、C 三谱线：0.486、0.588、0.656 μm"),
        ("MTF", "100 lp/mm 处 MTF > 0.4", "100 lp/mm 处约大于 0.55，满足要求"),
        ("结构形式", "库克三片式", "保持正-负-正三片结构，材料为 SK16/F2/SK16"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            set_cell_text(cell, text, size=11)
    add_para("由表中可见，系统主要设计指标均得到满足。优化后的系统在保持结构简单的基础上，实现了焦距、孔径、视场和成像质量之间的平衡。")

    add_heading("五、总结")
    add_para(
        "本次设计以 Zemax OpticStudio 为工具，对库克三片式镜头进行了从初始结构到优化结果的完整设计过程。"
        "首先根据 EFFL = 10 mm、相对孔径 1/2.8、2ω = 30° 和 100 lp/mm 处 MTF > 0.4 的要求完成系统设置；"
        "然后建立以 MTF/对比度为目标的评价函数，并加入焦距、厚度和空气间隔等边界约束；最后采用阻尼最小二乘法进行多轮优化，"
        "并通过布局图、点列图和 MTF 曲线评价结果。"
    )
    add_para(
        "从最终结果看，优化后的库克三片式镜头在 100 lp/mm 处的 MTF 高于 0.4，满足课程设计指标；"
        "三视场点列图显示系统具备较好的聚焦能力，结构上仍保持三片式镜头的简洁性。通过本次设计，可以进一步理解光学系统设计中"
        "“结构选择、变量控制、评价函数建立和像质分析”之间的关系，也认识到实际优化需要在像质、焦距、孔径、视场和加工可行性之间进行综合权衡。"
    )

    add_heading("附：优化过程截图汇总")
    add_para("以下截图来自本次 Zemax 优化过程，用于记录从系统设置、初始结构、变量调整到最终结果分析的主要步骤。")
    for img, cap in [
        ("step_06.png", "图11 第一轮优化运行窗口"),
        ("step_07.png", "图12 第一轮优化后的镜头数据"),
        ("step_08.png", "图13 中间优化结构参数"),
        ("step_09.png", "图14 中间轮次优化窗口"),
        ("step_11.png", "图15 中间结构布局图"),
    ]:
        add_picture(img, cap, width=5.4)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run()
        font_run(run, size=10.5)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 1
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
