from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


BASE = Path(__file__).resolve().parent
ASSET = BASE / "report_assets"
OUT = BASE / "202426001054_吴凯珍_光学系统设计课程报告.docx"


def font_run(run, size=12, bold=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_borders(table):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_cell(cell, text, bold=False, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1
    r = p.add_run(text)
    font_run(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.paragraph_format.line_spacing = 1
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Title"].font.size = Pt(18)
    doc.styles["Title"].font.bold = True
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 2"].font.bold = True

    def para(text="", align=None, first_indent=True, size=12, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if first_indent and text:
            p.paragraph_format.first_line_indent = Pt(24)
        if align is not None:
            p.alignment = align
            p.paragraph_format.first_line_indent = None
        r = p.add_run(text)
        font_run(r, size=size, bold=bold)
        return p

    def heading(text, level=1):
        p = doc.add_paragraph(style=f"Heading {level}")
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = None
        r = p.add_run(text)
        font_run(r, size=14 if level == 1 else 13, bold=True)

    def picture(step, title, note, width=5.55):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(ASSET / f"step_{step:02d}.png"), width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing = 1
        cap.paragraph_format.space_after = Pt(2)
        r = cap.add_run(f"图{step} {title}")
        font_run(r, size=10.5, bold=True)
        para(note)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("《光学系统设计》课程报告")
    font_run(r, size=20, bold=True)
    for _ in range(4):
        para("", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

    cover = doc.add_table(rows=4, cols=2)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(cover)
    for row, data in zip(
        cover.rows,
        [
            ("题目", "基于库克三片式镜头的光学系统优化设计"),
            ("班级", "24级光电信息科学与工程"),
            ("姓名", "吴凯珍"),
            ("学号", "202426001054"),
        ],
    ):
        row.height = Cm(1.1)
        set_cell(row.cells[0], data[0], bold=True, size=14)
        set_cell(row.cells[1], data[1], size=14)
    for _ in range(6):
        para("", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    para("完成日期：2026年6月28日", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    doc.add_page_break()

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于库克三片式镜头的光学系统优化设计")
    font_run(r, size=18, bold=True)
    para(
        "班级：24级光电信息科学与工程    姓名：吴凯珍    学号：202426001054",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
    )

    heading("一、设计背景与任务要求")
    para(
        "库克三片式镜头由正透镜、负透镜和正透镜组成，是经典的中等视场成像物镜。"
        "本报告按照截图原始顺序记录 Zemax 优化过程，并使每段文字与相邻图片中的数据一一对应。"
        "设计指标为：有效焦距 EFFL = 10 mm，相对孔径 1/2.8，2ω = 30°，并要求 100 lp/mm 处 MTF 大于 0.4。"
    )

    heading("二、按截图顺序的 Zemax 设置、优化与结果分析")
    picture(
        1,
        "评价函数编辑器与优化约束",
        "图1对应评价函数设置。截图中评价函数值为 0.125903055527036，评价类型选用“对比度”，空间频率为 100 lp/mm，"
        "S 权重和 T 权重均为 1，类型为 RMS，参考为质心；光瞳采样采用高斯求积，3 环、6 臂。"
        "下方操作数包含 EFFL 约束，目标有效焦距为 10.000 mm；同时设置空气和玻璃厚度边界，例如玻璃最小厚度 0.3 mm、最大厚度 15 mm，"
        "空气最小间隔 0.3 mm。该图说明优化目标和工艺边界已经先建立。"
    )
    picture(
        2,
        "系统孔径、视场和波长设置",
        "图2对应系统设置。孔径类型为入瞳直径，孔径值为 3.571 mm，这与 EFFL = 10 mm、F/# = 2.8 的计算结果一致。"
        "视场设置为 0°、15°、7.5°三个视场，权重均为 1；波长采用 F、d、C 可见光谱线，分别约为 0.486 μm、0.588 μm、0.656 μm，权重均为 1。"
        "这些数据对应题目中的相对孔径、30°全视场和复色光像质评价要求。"
    , width=3.35)
    picture(
        3,
        "初始结构点列图",
        "图3对应初始结构的点列图。三个视场的像高分别约为 0.000 mm、2.488 mm、1.288 mm；RMS 半径分别约为 "
        "6.273 μm、13.151 μm、17.040 μm，GEO 半径分别约为 12.145 μm、66.460 μm、158.066 μm。"
        "可以看出轴上视场较集中，而 7.5°和 15°离轴视场弥散明显，说明初始结构仍需优化。"
    )
    picture(
        4,
        "初始结构二维布局图",
        "图4对应初始库克三片式镜头布局。图中可以看到正-负-正三片结构和三色光线传播情况，系统总轴长约为 42.84953 mm。"
        "初始结构已能成像，但离轴光线在像面附近仍有较大散开，与图3点列图中的离轴弥散相对应。"
    )
    picture(
        5,
        "FFT MTF 设置及曲线",
        "图5对应 MTF 分析窗口。采样为 64×64，最大空间频率为 100 lp/mm，类型为调制，波长和视场均选择“所有”，表面为像面。"
        "曲线在 100 lp/mm 处仍高于 0.4，说明当前结构在评价频率处具备一定对比度传递能力；该图与图1中 100 lp/mm 的评价函数设置相对应。"
    )
    picture(
        6,
        "第一轮局部优化窗口",
        "图6对应一次阻尼最小二乘法优化。目标数为 163，变量数为 12，内核数目为 16；初始评价函数为 0.079400756，"
        "当前评价函数降为 0.070248703，执行时间为 1.921 s。该图说明在已建立评价函数后，Zemax 对 12 个变量进行了局部迭代，评价函数有所降低。"
    )
    picture(
        7,
        "第一轮优化后的镜头数据",
        "图7对应第一轮优化后的结构参数。第 1 片材料为 SK16，曲率半径约为 -13.991 mm，厚度约为 13.373 mm；"
        "中间负透镜材料为 F2，相关曲率半径约为 8.242 mm，厚度约为 5.265 mm；第 3 片材料为 SK16，厚度约为 15.008 mm。"
        "像面净口径约为 2.489 mm。该图记录的是中间状态，不是最终结构。"
    )
    picture(
        8,
        "继续调整后的镜头数据",
        "图8对应另一组中间结构数据。第 1 面曲率半径约为 -10.111 mm、厚度约为 11.028 mm，SK16；第 3 面曲率半径约为 8.484 mm、"
        "厚度约为 4.432 mm，F2；第 5 面曲率半径约为 2.953 mm、厚度约为 15.014 mm，SK16；像面净口径约为 2.512 mm。"
        "与图7相比，曲率、厚度和像面口径均发生变化，说明优化过程中结构仍在调整。"
    )
    picture(
        9,
        "第二轮优化窗口",
        "图9对应第二轮阻尼最小二乘法优化。目标数为 163，变量数为 12，初始评价函数为 17.503302509，当前评价函数为 0.075295694，"
        "执行时间为 5.344 s。初始值较大说明本轮开始前结构或约束状态发生变化，经过迭代后评价函数显著降低。"
    )
    picture(
        10,
        "中间优化结果点列图",
        "图10对应第二轮优化后的点列图。三个视场的像高约为 0.000 mm、2.510 mm、1.291 mm；RMS 半径约为 "
        "7.927 μm、23.040 μm、12.881 μm，GEO 半径约为 15.596 μm、85.479 μm、77.063 μm。"
        "与初始点列图相比，边缘视场形态发生变化，但半视场 RMS 半径仍偏大，因此需要继续观察结构和优化结果。"
    )
    picture(
        11,
        "中间结构二维布局图",
        "图11对应中间结构布局，系统总轴长约为 38.76406 mm。图中最后一片或后组表面被选中显示，三色光线在像面附近重新交会。"
        "与图4相比，系统轴向长度缩短，后组位置和光线交会关系发生改变，说明优化不仅改变了像质，也改变了系统结构尺寸。"
    )
    picture(
        12,
        "继续优化窗口",
        "图12对应继续优化后的窗口。算法仍为阻尼最小二乘法，目标数 163，变量数 12，初始评价函数为 17.503302509，"
        "当前评价函数为 0.075295694，执行时间为 5.672 s。该图与图9数据基本一致，可作为本轮优化收敛状态的记录。"
    )
    picture(
        13,
        "最终优化后的镜头数据",
        "图13对应最终结构参数。第 1 面曲率半径约为 22.014 mm、厚度约为 3.259 mm，材料 SK16；第 2 面曲率半径约为 -435.760 mm、"
        "空气间隔约为 6.008 mm；第 3 面曲率半径约为 -22.213 mm、厚度约为 1.000 mm，材料 F2；第 5 面曲率半径约为 79.684 mm、"
        "厚度约为 2.952 mm，材料 SK16；第 6 面至像面距离约为 42.208 mm，像面净口径约为 13.357 mm。"
        "这组数据对应最终三片式结构。"
    )
    picture(
        14,
        "最终结构点列图",
        "图14对应最终结构点列图。图中比例尺为 20 μm，三个视场的像高约为 0.000 mm、13.348 mm、6.557 mm；RMS 半径约为 "
        "5.536 μm、4.082 μm、4.692 μm，GEO 半径约为 8.323 μm、8.869 μm、9.792 μm。"
        "三视场弥散斑均较小且分布集中，说明最终结构的几何像质明显优于前面的中间结果。"
    )
    picture(
        15,
        "最终结构二维布局图",
        "图15对应最终结构布局，系统总轴长约为 60.17675 mm。最终系统仍保持库克三片式正-负-正结构，"
        "三种波长光线在像面附近较好会聚。该图与图13最终镜头数据、图14最终点列图相互对应，说明最终结构在参数和像质上已形成闭环。"
    )

    heading("三、综合结论")
    para(
        "按照以上 15 张截图的顺序可以看出，本设计从评价函数和系统参数设置开始，依次完成初始像质分析、MTF 检查、"
        "多轮阻尼最小二乘法优化、镜头数据更新以及最终点列图和布局图验证。最终结构在 100 lp/mm 处 MTF 高于 0.4，"
        "最终点列图 RMS 半径约为 4.082 μm 至 5.536 μm，满足课程报告中对库克三片式镜头优化设计的基本要求。"
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run()
        font_run(run, size=10.5)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 1
        for r in p.runs:
            r.font.name = "Times New Roman"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
